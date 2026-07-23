"""
Fill the ACIP (华进) invention disclosure .docx template with content.

CURRENT SCOPE: ACIP-only. The TEMPLATES dict below currently registers
only the `acip` template, and the 16 field coordinates (row, col) are
hard-coded against ACIP's specific .docx table layout. Despite the
config-driven shape (TEMPLATES dict + --template flag), this script is
NOT yet a multi-agency tool — adding another agency requires deriving
new (row, col) coordinates and field names for that template.

The `inspect` subcommand IS generic: it prints any .docx's table layout
(merged-cell aware) to help derive coordinates when onboarding a new
agency. See SKILL.md "Adding a new agency template" for the workflow.

Usage
-----
1. Fill the ACIP template with content from JSON:
     python fill_acip_template.py fill \
         --template acip \
         --content invention.json \
         --output "Disclosure-ACIP-ARGesture-20260720.docx"

   If --content is omitted, a built-in sample (AR gesture) is used for testing.

2. Inspect any .docx template's table structure (for onboarding new agencies):
     python fill_acip_template.py inspect \
         --docx "path/to/new_agency_template.docx"

   This prints each table's row/column layout with merged-cell detection,
   so you can derive the field->(row,col) mapping for a new template config.

3. List registered templates:
     python fill_acip_template.py list

Author: patent-forge skill
"""
from __future__ import annotations

import argparse
import json
import sys
import pathlib
from dataclasses import dataclass, field
from typing import Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ===========================================================================
# Template registry — programmatic mirror of assets/templates/template_registry.md
# ===========================================================================
# Each template defines:
#   docx_path : relative path (from skill root) to the original .docx template
#   table_idx : index of the main table in doc.tables (usually 0 or 1)
#   fields    : field_name -> (row, col) where content should be written
#
# CRITICAL: when adding a new template, run `inspect` first to find the
# correct (row, col) for each field. Never guess — merged cells in
# python-docx make col index misleading (a merged cell appears at multiple
# col positions pointing to the same _tc).

TEMPLATES: Dict[str, dict] = {
    "acip": {
        "docx_path": "assets/raw_templates/acip_invention_disclosure.docx",
        "table_idx": 1,   # ACIP puts the instructions in table 0, main form in table 1
        "fields": {
            # Header rows
            "case_name":    (1, 1),   # 专利申请案件名称 -> value cell
            "inventors":    (2, 1),
            "applicant":    (3, 1),
            "tech_contact": (3, 4),   # same row, right-side cell
            "tech_phone":   (4, 1),
            "tech_email":   (4, 4),
            "published":    (5, 4),   # 是/否
            # Body sections — rows 7/8 keep question label in cell 0, answer in cell 2
            "bg_field":      (7,  2),
            "closest_prior": (8,  2),
            # Other sections are full-width merged cells
            "problems":      (10, 0),
            "invention_pts": (12, 0),
            "details":       (14, 0),
            "effects":       (16, 0),
            "alternatives":  (18, 0),
            "terminology":   (20, 0),
            "references":    (22, 0),
        },
    },
    # Future templates (e.g. other agencies) go here. Use `inspect` to derive.
}


# ===========================================================================
# Cell helpers
# ===========================================================================
def clear_cell(cell) -> None:
    """Remove all <w:p> in a cell, then add one empty paragraph.
    Preserves cell properties (tcPr: width, borders, merge)."""
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    p = OxmlElement("w:p")
    tc.append(p)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Replace cell content with multi-line text. Cell formatting inherited."""
    clear_cell(cell)
    p = cell.paragraphs[0]
    lines = text.split("\n") if isinstance(text, str) else list(text)
    for li, line in enumerate(lines):
        if li > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        run.bold = bold


# ===========================================================================
# Fill logic
# ===========================================================================
@dataclass
class FillResult:
    filled_fields: List[str] = field(default_factory=list)
    skipped_fields: List[str] = field(default_factory=list)
    output_path: str = ""


def fill_template(
    template_id: str,
    content: Dict[str, str],
    output_path: str,
    skill_root: Optional[pathlib.Path] = None,
) -> FillResult:
    """Open template .docx, fill cells per TEMPLATES[template_id], save."""
    if template_id not in TEMPLATES:
        raise ValueError(
            f"Unknown template '{template_id}'. Registered: {list(TEMPLATES)}"
        )
    cfg = TEMPLATES[template_id]

    # Resolve docx path
    if skill_root is None:
        skill_root = _default_skill_root()
    docx_path = (pathlib.Path(skill_root) / cfg["docx_path"])
    if not docx_path.exists():
        raise FileNotFoundError(f"Template .docx not found: {docx_path}")

    doc = Document(str(docx_path))
    if len(doc.tables) <= cfg["table_idx"]:
        raise RuntimeError(
            f"Template has {len(doc.tables)} tables, expected index "
            f"{cfg['table_idx']}. Run `inspect` to verify."
        )
    table = doc.tables[cfg["table_idx"]]

    result = FillResult(output_path=output_path)
    for field_name, (row, col) in cfg["fields"].items():
        if field_name not in content:
            result.skipped_fields.append(field_name)
            continue
        value = content[field_name]
        # Special handling for `published` boolean-like
        if field_name == "published":
            value = "是" if str(value).lower() in ("yes", "true", "1", "是") else "否"
        if row >= len(table.rows):
            result.skipped_fields.append(f"{field_name} (row {row} out of range)")
            continue
        cells = table.rows[row].cells
        if col >= len(cells):
            result.skipped_fields.append(f"{field_name} (col {col} out of range)")
            continue
        set_cell_text(cells[col], str(value))
        result.filled_fields.append(field_name)

    doc.save(output_path)
    return result


def _default_skill_root() -> pathlib.Path:
    # scripts/ is one level below skill root
    return pathlib.Path(__file__).resolve().parent.parent


# ===========================================================================
# Inspect logic — for adding new templates
# ===========================================================================
def inspect_template(docx_path: str) -> None:
    """Print every table's structure to help derive field->(row,col) mapping."""
    doc = Document(docx_path)
    print(f"File: {docx_path}")
    print(f"Tables: {len(doc.tables)}")
    for ti, t in enumerate(doc.tables):
        print(f"\n=== Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols ===")
        for ri, row in enumerate(t.rows):
            seen_tc = []
            cells_info = []
            for ci, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    tag = "<merged>"
                else:
                    seen_tc.append(tc_id)
                    txt = cell.text.strip().replace("\n", " | ")
                    tag = (txt[:60] + "…") if len(txt) > 60 else txt
                cells_info.append(f"[{ci}]{tag!r}")
            print(f"  Row {ri:2d}: {' '.join(cells_info)}")


# ===========================================================================
# Built-in sample content (AR gesture) — used when --content is omitted
# ===========================================================================
SAMPLE_CONTENT = {
    "case_name":    "一种基于手势识别的增强现实交互方法及系统",
    "inventors":    "张三、李四",
    "applicant":    "XX科技有限公司",
    "tech_contact": "王五",
    "tech_phone":   "138-0000-0000",
    "tech_email":   "wangwu@example.com",
    "published":    "no",
    "bg_field": (
        "本发明涉及增强现实（Augmented Reality, AR）与人机交互技术领域，"
        "特别涉及一种基于手势识别的 AR 交互方法及系统。\n"
        "增强现实技术将虚拟信息叠加至真实世界，已在工业维修、教育培训、医疗辅助等领域得到广泛应用。"
        "手势作为一种自然、直观的交互方式，被认为是 AR 场景中最具潜力的输入模态之一。"
        "然而，现有 AR 手势交互方案存在以下不足：传统基于外观的方法在复杂光照和遮挡场景下识别精度显著下降；"
        "现有轻量化模型难以同时兼顾精度与实时性，端侧推理延迟普遍超过 80ms；"
        "手势指令映射逻辑僵化，缺乏基于上下文的自适应能力。"
    ),
    "closest_prior": (
        "经全球专利库（EPO DOCDB/INPADOC，覆盖 100+ 国家）检索，与本方案最接近的现有技术为：\n"
        "（1）CN114440000A（2022）——一种基于手势的 AR 标注方法。该专利通过肤色分割与模板匹配识别预设手势，成本低但精度受光照影响大。\n"
        "（2）US20210373015A1（2021）——Wearable AR gesture control system。采用 MediaPipe 端侧推理，延迟约 50ms。\n"
        "注：以\"AR gesture recognition real-time end-side\"为关键词在全球专利库中未检索到直接相关专利。"
    ),
    "problems": (
        "现有技术存在以下不足：\n"
        "（1）缺乏一种在端侧设备上同时兼顾高精度（>95%）与低延迟（<30ms）的手部关键点检测网络结构；\n"
        "（2）缺乏一种基于时序卷积网络（TCN）的动态手势序列建模方法；\n"
        "（3）缺乏一种基于场景上下文的自适应手势指令映射机制。"
    ),
    "invention_pts": (
        "本发明提出一种基于手势识别的 AR 交互方法及系统，核心技术方案包括三个层面：\n"
        "（1）关键点检测层面：提出基于 MobileNetV3 主干并经通道剪枝的轻量化手部关键点检测网络。\n"
        "（2）动态手势识别层面：提出基于时序卷积网络（TCN）的序列建模方法。\n"
        "（3）指令映射层面：提出基于场景上下文的自适应映射策略。"
    ),
    "details": (
        "4.1 系统整体架构\n本方案所述 AR 交互系统包括：图像采集模块、关键点检测模块、手势识别模块、指令映射模块、AR 渲染模块。\n"
        "【图 1】基于手势识别的 AR 交互系统整体架构示意图\n\n"
        "4.2 轻量化手部关键点检测网络\n主干采用 MobileNetV3-Large，输入 224×224×3，输出 21 个手部关键点坐标。通道剪枝比例 0.5 ≤ r ≤ 0.75。\n"
        "【图 2】轻量化手部关键点检测网络结构示意图\n\n"
        "4.3 基于 TCN 的动态手势识别\n步骤一：从连续 T=30 帧关键点序列中提取时空特征；\n步骤二：4 层时序卷积（kernel=3，dilation=[1,2,4,8]）建模长程依赖；\n步骤三：经全连接层输出 N 类动态手势概率分布。\n"
        "【图 3】动态手势识别方法流程图\n\n"
        "4.4 基于场景上下文的自适应指令映射\n指令映射模块维护\"空闲/选择/操作/确认\"四态状态机，同一手势在不同状态下映射为不同交互指令。\n"
        "【图 4】指令自适应映射机制示意图"
    ),
    "effects": (
        "（1）通过轻量化手部关键点检测网络，在保证识别精度≥95% 的前提下，将端侧推理延迟降至 30ms 以内。\n"
        "（2）通过基于 TCN 的动态手势序列建模，相比传统 3D CNN 计算量降低约 60%。\n"
        "（3）通过场景上下文自适应映射机制，使同一手势在不同场景下产生不同交互效果，显著提升了交互自然性。"
    ),
    "alternatives": (
        "替代方案一：采用 Transformer 替代 TCN 进行动态手势建模，精度提升约 2%，但计算量增加 20%。\n"
        "替代方案二：采用 RGB-D 融合输入替代纯深度图输入，弱光环境下精度提升 3-5%。"
    ),
    "terminology": (
        "术语/缩略语 | 解释说明\n"
        "AR | Augmented Reality，增强现实\n"
        "CNN | Convolutional Neural Network，卷积神经网络\n"
        "TCN | Temporal Convolutional Network，时序卷积网络\n"
        "MediaPipe | Google 开源的跨平台 ML 推理框架，含手部关键点检测模型\n"
        "MobileNetV3 | Google 提出的轻量化卷积神经网络架构\n"
        "RGB-D | Red Green Blue-Depth，彩色图+深度图四通道输入"
    ),
    "references": (
        "[1] CN114440000A (2022) — 一种基于手势的 AR 标注方法。\n"
        "[2] US20210373015A1 (2021) — Wearable AR gesture control system。\n"
        "[3] Zhang et al., \"MobileNetV3: Searching for MobileNetV3\", ICCV 2019。\n"
        "[4] Bai et al., \"An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling\", arXiv 2018。"
    ),
}


# ===========================================================================
# CLI
# ===========================================================================
def main():
    parser = argparse.ArgumentParser(
        prog="fill_acip_template",
        description="Fill patent disclosure templates with content",
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    # fill
    p_fill = sub.add_parser("fill", help="Fill a registered template")
    p_fill.add_argument("--template", default="acip", choices=list(TEMPLATES))
    p_fill.add_argument("--content", help="Path to JSON file with field values")
    p_fill.add_argument("--output", required=True, help="Output .docx path")
    p_fill.add_argument("--skill-root", help="Override skill root directory")

    # inspect
    p_ins = sub.add_parser("inspect", help="Inspect a template's table layout")
    p_ins.add_argument("--docx", required=True, help="Template .docx to inspect")

    # list
    sub.add_parser("list", help="List registered templates")

    args = parser.parse_args()

    if args.cmd == "list":
        print("Registered templates:")
        for tid, cfg in TEMPLATES.items():
            n_fields = len(cfg["fields"])
            print(f"  {tid:10s}  docx={cfg['docx_path']}  "
                  f"table_idx={cfg['table_idx']}  fields={n_fields}")
        return

    if args.cmd == "inspect":
        inspect_template(args.docx)
        return

    if args.cmd == "fill":
        content = SAMPLE_CONTENT
        if args.content:
            with open(args.content, "r", encoding="utf-8") as f:
                content = json.load(f)
        result = fill_template(args.template, content, args.output, args.skill_root)
        print(f"[OK] Saved: {result.output_path}")
        print(f"  Filled ({len(result.filled_fields)}): {result.filled_fields}")
        if result.skipped_fields:
            print(f"  Skipped ({len(result.skipped_fields)}): {result.skipped_fields}")


if __name__ == "__main__":
    main()
