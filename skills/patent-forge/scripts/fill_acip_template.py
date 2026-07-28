"""
Fill the ACIP (华进) invention disclosure .docx template with content.

CURRENT SCOPE: ACIP-only. The TEMPLATES dict below currently registers
only the `acip` template, and the 16 field coordinates (row, col) are
hard-coded against ACIP's specific .docx table layout. Despite the
config-driven shape (TEMPLATES dict + --template flag), this script is
NOT yet a multi-agency tool — adding another agency requires deriving
new (row, col) coordinates and field names for that template.

The `inspect` subcommand IS generic: it prints any .docx's table layout
(merged-cell aware) to help derive coordinates when onboarding a new
agency. See SKILL.md "Adding a new agency template" for the workflow.

Usage
-----
1. Fill the ACIP template with content from JSON (figures embedded by default):
     python fill_acip_template.py fill \
         --template acip \
         --content invention.json \
         --output "Disclosure-ACIP-ARGesture-20260720.docx"

   Figures are auto-discovered from `<skill_root>/../04-diagrams/` (the
   standard patent-forge Phase 3 output dir) or from `--figures-dir` if
   passed. SVG companions are skipped (Word cannot embed SVG inline); PNG is
   always preferred. Pass `--no-figures` to disable auto-discovery.

   If --content is omitted, a built-in sample (AR gesture) is used for testing.

2. Inspect any .docx template's table structure (for onboarding new agencies):
     python fill_acip_template.py inspect \
         --docx "path/to/new_agency_template.docx"

   This prints each table's row/column layout with merged-cell detection,
   so you can derive the field->(row,col) mapping for a new template config.

3. List registered templates:
     python fill_acip_template.py list

Author: patent-forge skill
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
import pathlib
from dataclasses import dataclass, field
from typing import Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ===========================================================================
# Template registry — programmatic mirror of assets/templates/template_registry.md
# ===========================================================================
# Each template defines:
#   docx_path : relative path (from skill root) to the original .docx template
#   table_idx : index of the main table in doc.tables (usually 0 or 1)
#   fields    : field_name -> (row, col) where content should be written
#
# CRITICAL: when adding a new template, run `inspect` first to find the
# correct (row, col) for each field. Never guess — merged cells in
# python-docx make col index misleading (a merged cell appears at multiple
# col positions pointing to the same _tc).

TEMPLATES: Dict[str, dict] = {
    "acip": {
        "docx_path": "assets/raw_templates/acip_invention_disclosure.docx",
        "table_idx": 1,   # ACIP puts the instructions in table 0, main form in table 1
        "fields": {
            # Header rows
            "case_name":    (1, 1),   # 专利申请案件名称 -> value cell
            "inventors":    (2, 1),
            "applicant":    (3, 1),
            "tech_contact": (3, 4),   # same row, right-side cell
            "tech_phone":   (4, 1),
            "tech_email":   (4, 4),
            "published":    (5, 4),   # 是/否
            # Body sections — rows 7/8 keep question label in cell 0, answer in cell 2
            "bg_field":      (7,  2),
            "closest_prior": (8,  2),
            # Other sections are full-width merged cells
            "problems":      (10, 0),
            "invention_pts": (12, 0),
            "details":       (14, 0),
            "effects":       (16, 0),
            "alternatives":  (18, 0),
            "terminology":   (20, 0),
            "references":    (22, 0),
        },
    },
    # Future templates (e.g. other agencies) go here. Use `inspect` to derive.
}


# ===========================================================================
# Cell helpers
# ===========================================================================
def clear_cell(cell) -> None:
    """Remove all paragraphs AND nested tables in a cell, then add one empty
    paragraph. Preserves cell properties (tcPr: width, borders, merge).

    Without removing <w:tbl> children, pre-existing sample tables shipped in a
    template (e.g. ACIP's terminology table) would survive and accumulate
    alongside newly-rendered content.
    """
    tc = cell._tc
    # Remove both <w:p> (paragraphs) and <w:tbl> (nested tables) in any order.
    for tag in ("w:p", "w:tbl"):
        for el in tc.findall(qn(tag)):
            tc.remove(el)
    p = OxmlElement("w:p")
    tc.append(p)


# ===========================================================================
# Emoji / status-symbol stripping
# ===========================================================================
# Patent disclosure documents are formal legal-adjacent artifacts. Internal
# workflow markers (🔴🟠🟡🟢⚠️✅❌⛔🛑⚡🔒 and similar) are useful during
# AI-assisted drafting but must NOT appear in the final .docx/.md delivered
# to a patent agent. This regex matches the common emoji ranges + an
# explicit allowlist of status glyphs seen across the patent-forge skills.

_EMOJI_RE = re.compile(
    "["
    "\U0001F300-\U0001FAFF"   # symbols & pictographs (🔴🟠🟡🟢✅❌⛔🛑⚡🔒 ...)
    "\U00002600-\U000027BF"   # misc symbols (☀ ☂ ☑ ☒ ⚠ ...)
    "\U0001F000-\U0001F2FF"   # mahjong / dominoes / cards (rare here, safe to drop)
    "\U00002B00-\U00002BFF"   # arrows/stars (➜ ⭐) — note: → (U+2192) is OUTSIDE this range and preserved
    "\u200D"                  # ZWJ used in emoji composition
    "\uFE0F"                  # VS-16 (emoji variation selector, e.g. trailing ⚠️)
    "\u20E3"                  # combining enclosing keycap
    "]+",
    flags=re.UNICODE,
)


def _strip_status_symbols(text: str) -> str:
    """Remove emoji / status glyphs from text, collapsing the resulting
    double spaces (e.g. "（🔴 最接近）" -> "（最接近）", "✅ 优秀" -> "优秀").
    Non-emoji punctuation (Chinese/ASCII brackets, arrows like → used as
    flow notation in method steps) is preserved.
    """
    if not isinstance(text, str):
        return text
    cleaned = _EMOJI_RE.sub("", text)
    # Collapse HORIZONTAL spaces left behind where an emoji sat between CJK
    # chars or before/after a bracket: "（  text" / "text  ）" / "A  B".
    # IMPORTANT: must NOT match \n / \r — paragraph breaks (\n\n) in the
    # content must survive so set_cell_text can render them as <w:br/>.
    # \s would also eat \n, so use [ \t\f\v] explicitly.
    cleaned = re.sub(r"[ \t\f\v]{2,}", " ", cleaned)
    cleaned = re.sub(r"([（(])\s+", r"\1", cleaned)   # "（ text" -> "（text"
    cleaned = re.sub(r"\s+([）)])", r"\1", cleaned)   # "text ）" -> "text）"
    return cleaned.strip()


# ===========================================================================
# Structured-input coercion (list -> markdown)
# ===========================================================================
# Callers (and content JSON authored by the AI) may legitimately pass a
# Python list for fields that are inherently tabular (terminology,
# references) — e.g. `terminology: [["术语","英文","中文"], ["LLM","...","..."]]`.
# The fill loop does `str(value)` which turns such a list into its Python
# repr (a single-line literal "[['术语','英文',...], ...]") that the markdown
# table parser cannot recognize. This helper converts:
#   - list[list[str]]  -> markdown pipe table (header row + separator + rows)
#   - list[str]        -> markdown bullet list (one "- item" per line)
#   - list[dict]       -> markdown pipe table built from dict keys (cols)
#                          in first-dict insertion order
# Non-list / non-iterable input is returned unchanged. Strings get a fast
# pass-through (so existing markdown-text callers are unaffected).

def _md_escape_cell(s: str) -> str:
    """Escape pipe and newline inside a markdown table cell so the row stays
    on one line. Backslash-escape per CommonMark; literal \n in a cell is
    replaced with ' / ' to keep the row flat (Word cell wraps anyway)."""
    if s is None:
        return ""
    s = str(s).replace("\\", "\\\\").replace("|", "\\|").replace("\n", " / ")
    return s


def _coerce_structured_to_markdown(value) -> str:
    """Convert list/dict structured content into markdown text. Strings and
    other scalars pass through unchanged. Used at the top of set_cell_text
    so the existing _parse_md_tables + paragraph renderers handle structured
    JSON the same way they handle hand-written markdown strings.
    """
    if isinstance(value, str):
        return value
    # list[list[str]] -> markdown table
    if isinstance(value, list) and value and all(isinstance(r, (list, tuple)) for r in value):
        rows = [[_md_escape_cell(c) for c in r] for r in value]
        n_cols = max(len(r) for r in rows)
        # Pad ragged rows to n_cols
        rows = [r + [""] * (n_cols - len(r)) for r in rows]
        header = rows[0]
        body = rows[1:]
        sep = "|".join(["---"] * n_cols)
        lines = ["| " + " | ".join(header) + " |", "| " + sep + " |"]
        for r in body:
            lines.append("| " + " | ".join(r) + " |")
        return "\n".join(lines)
    # list[dict] -> markdown table (cols = first dict's keys, insertion order)
    if isinstance(value, list) and value and all(isinstance(r, dict) for r in value):
        # Collect union of keys preserving first-seen order across all rows
        keys: List[str] = []
        for r in value:
            for k in r.keys():
                if k not in keys:
                    keys.append(k)
        rows = [[_md_escape_cell(r.get(k, "")) for k in keys] for r in value]
        sep = "|".join(["---"] * len(keys))
        lines = ["| " + " | ".join(keys) + " |", "| " + sep + " |"]
        for r in rows:
            lines.append("| " + " | ".join(r) + " |")
        return "\n".join(lines)
    # list[str] -> markdown bullet list
    if isinstance(value, list):
        return "\n".join("- " + _md_escape_cell(s) for s in value)
    # dict -> markdown bullet list (key: value)
    if isinstance(value, dict):
        return "\n".join(f"- **{k}**: {_md_escape_cell(v)}" for k, v in value.items())
    # Fallback: stringify (numbers, bools, None)
    return str(value)


# ===========================================================================
# Template boilerplate hint stripping
# ===========================================================================
# The ACIP (and similar agency) .docx templates ship each answer cell with an
# instruction hint, e.g. row 20 (terminology) starts with
#   "请对本交底书中提及的关键术语、技术缩略语进行解释说明..."
# and row 22 (references) starts with
#   "（对于理解交底书中的技术方案有帮助的专利/论文/期刊，如有则填写）"
# When the AI generates content it tends to echo these same hint phrases as a
# prefix to the real answer, producing duplicated / leaked boilerplate in the
# final doc. These prefixes are stripped from the START of any field value so
# only the actual content remains. Matching is prefix-anchored and tolerant of
# full-width/half-width parentheses and CJK punctuation variants.

_BOILERPLATE_HINT_PREFIXES = (
    # Terminology section hint (ACIP row 20)
    "请对本交底书中提及的关键术语、技术缩略语进行解释说明",
    "请对本交底书中提及的关键术语、技术缩略语进行解释",
    "请对本交底书",
    # References section hint (ACIP row 22) — appears with full-width or
    # half-width parentheses, with or without trailing "，如有则填写"
    "（对于理解交底书中的技术方案有帮助的专利/论文/期刊",
    "(对于理解交底书中的技术方案有帮助的专利/论文/期刊",
    "（对于理解交底书中的技术方案有帮助的专利",
    "(对于理解交底书中的技术方案有帮助的专利",
    "对于理解交底书中的技术方案有帮助的专利/论文/期刊",
)


def _strip_boilerplate_hints(text: str) -> str:
    """Remove template boilerplate instruction hints from the start of `text`.

    Handles the common variants: hint followed by '：' then content, hint
    wrapped in parentheses on its own line, hint followed by newlines. Keeps
    the remainder intact. Idempotent: if no hint is present, returns text as-is.
    """
    if not isinstance(text, str) or not text:
        return text
    for hint in _BOILERPLATE_HINT_PREFIXES:
        if text.startswith(hint):
            rest = text[len(hint):]
            # Drop a trailing instruction clause up to and including the first
            # '：' / ':' / '）' / ')' / newline so we keep only real content.
            # Examples:
            #   "...解释说明：\n\n| 术语 ..."  -> keep "| 术语 ..."
            #   "...解释说明，如果有英文缩写，必须给出...）\n\n" -> keep ""
            #   "（对于...如有则填写）\n\n[1] ..." -> keep "[1] ..."
            rest = re.sub(
                r"^[：:）)\s]*", "", rest, count=1
            )
            # If the hint itself carried a full instruction sentence (e.g. the
            # long terminology hint with "如果有英文缩写，必须给出..."), strip
            # everything up to the first markdown table header or reference
            # marker or the first newline-run that introduces real content.
            # Heuristic: cut any leading prose that does NOT start with a
            # table pipe '|', a reference '[N]', or another structural token.
            rest = re.sub(
                r"\A(?:[^|\[\n]*?)(?=\||\[|\n\n|\Z)", "", rest, count=1,
                flags=re.DOTALL,
            )
            text = rest.lstrip(" \t")
            break
    # Collapse any leading blank lines left behind.
    return text.lstrip("\n").strip()


# ===========================================================================
# Markdown table parsing & rendering
# ===========================================================================
# A markdown pipe-table block is detected as 2+ consecutive lines where:
#   - line 0 contains a pipe `|`
#   - line 1 is EITHER:
#       (a) a separator row like `|---|---|` (only dashes/colons/pipes/spaces), OR
#       (b) another pipe-data row (tolerant mode: a missing separator is
#           synthesized so LLM/human-generated markdown tables without the
#           `|---|---|` line still render as native Word tables).
# Captured tables are rendered as native Word nested tables instead of text.

_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def _is_md_table_line(line: str) -> bool:
    return "|" in line


def _is_md_table_sep(line: str) -> bool:
    return bool(_MD_TABLE_SEP_RE.match(line))


def _split_md_row(line: str) -> List[str]:
    """Split a markdown table row into trimmed cells, stripping outer pipes."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


@dataclass
class _MdTable:
    header: List[str]
    rows: List[List[str]]


def _parse_md_tables(text: str) -> List:
    """Segment text into a list of blocks: ('text', str) or ('table', _MdTable).

    Keeps plain text blocks intact (preserving their newlines) so the caller
    can render them as paragraphs; extracted markdown tables become _MdTable
    objects to be rendered as native Word tables.

    Tolerant mode: when a pipe-header line is immediately followed by another
    pipe-data line (no `|---|---|` separator), the separator is synthesized
    and the table is still parsed. This handles the common LLM/human failure
    mode of omitting the separator row.
    """
    lines = text.split("\n") if isinstance(text, str) else list(text)
    blocks: List = []
    i = 0
    n = len(lines)
    text_buf: List[str] = []

    def flush_text() -> None:
        if text_buf:
            blocks.append(("text", "\n".join(text_buf)))
            text_buf.clear()

    while i < n:
        line = lines[i]
        next_is_sep = i + 1 < n and _is_md_table_sep(lines[i + 1])
        # Tolerant mode: header followed directly by another pipe-data line
        # (no separator). Treat line[i+1] as the first data row.
        next_is_data_no_sep = (
            not next_is_sep
            and i + 1 < n
            and _is_md_table_line(lines[i + 1])
            and not _is_md_table_sep(lines[i + 1])
        )
        # Detect table start: current line has a pipe AND next line is either
        # a separator OR another pipe-data row (tolerant mode).
        if _is_md_table_line(line) and (next_is_sep or next_is_data_no_sep):
            flush_text()
            header = _split_md_row(line)
            i += 1  # skip header (separator, if present, is consumed below)
            # If a real separator exists, skip it; otherwise the synthesized
            # one needs no skip (i already points at first data row).
            if next_is_sep:
                i += 1  # skip separator
            rows: List[List[str]] = []
            while i < n and _is_md_table_line(lines[i]) and not _is_md_table_sep(lines[i]):
                rows.append(_split_md_row(lines[i]))
                i += 1
            blocks.append(("table", _MdTable(header=header, rows=rows)))
        else:
            text_buf.append(line)
            i += 1
    flush_text()
    return blocks


def _render_md_table(cell, md: _MdTable, bold: bool = False) -> None:
    """Append a native Word nested table inside `cell` for the markdown table."""
    n_cols = max([len(md.header)] + [len(r) for r in md.rows])
    n_rows = 1 + len(md.rows)  # header + data
    tbl = cell.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    def _write(w_row: int, w_col: int, value: str, is_header: bool) -> None:
        if w_col >= n_cols:
            return
        c = tbl.cell(w_row, w_col)
        # Clear the default empty paragraph then write text as a single run.
        clear_cell(c)
        para = c.paragraphs[0]
        # Header-row text is horizontally centered (固化在代码中)：表头如
        # "术语/缩略语" / "解释说明" 居中显示，数据行保持默认左对齐。
        if is_header:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(value)
        run.bold = is_header or bold
        run.font.size = Pt(9)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if is_header:
            # shade header row light blue for readability
            _shade_cell(c, "D9E2F3")

    # Header row
    for ci, val in enumerate(md.header):
        _write(0, ci, val, is_header=True)
    # Data rows
    for ri, row in enumerate(md.rows, start=1):
        for ci, val in enumerate(row):
            _write(ri, ci, val, is_header=False)


def _shade_cell(cell, hex_fill: str) -> None:
    """Apply a background fill color to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, allow_tables: bool = True) -> str:
    """Replace cell content with multi-line text.

    Markdown pipe-tables embedded in `text` are rendered as native Word nested
    tables (each with a header row + data rows and Table Grid borders). Plain
    text blocks between/around tables are rendered as paragraphs.
    Cell formatting is otherwise inherited.

    `allow_tables`: when False, any markdown pipe-table found in `text` is
    NOT rendered as a native Word table — the table lines are kept as plain
    text paragraphs (pipe characters preserved). Used by `fill_template` to
    enforce the "only terminology/references sections may contain tables"
    constraint in the disclosure document.

    Status/emoji glyphs are stripped before rendering (formal disclosure).
    Template boilerplate instruction hints (e.g. ACIP's "请对本交底书中提及的
    关键术语..." or "（对于理解交底书中的技术方案有帮助的专利/论文/期刊）")
    are also stripped from the start of the value so they do not leak into the
    final deliverable.
    Returns the cleaned text (for caller bookkeeping).
    """
    # Normalize structured input (list-of-lists / list-of-strings) to a
    # markdown representation so the existing table/bullet renderers handle
    # it. Without this, str(value) at the call site turns a Python list into
    # its repr (e.g. "[['术语', '英文', '中文'], ...]") which renders as one
    # ugly literal string instead of a native Word table.
    text = _coerce_structured_to_markdown(text)
    text = _strip_status_symbols(text)
    text = _strip_boilerplate_hints(text)
    clear_cell(cell)
    blocks = _parse_md_tables(text) if allow_tables else [("text", text)]
    # When allow_tables is False but the text contains table-looking blocks,
    # _parse_md_tables is skipped entirely; pipes stay as literal characters.
    first = True
    for kind, payload in blocks:
        if kind == "text":
            # Render each text block as its own paragraph (newlines -> breaks).
            lines = payload.split("\n") if payload else [""]
            # Drop a single trailing empty line that arises from table detection.
            if len(lines) > 1 and lines[-1] == "":
                lines = lines[:-1]
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            for li, line in enumerate(lines):
                if li > 0:
                    run = p.add_run()
                    run.add_break()
                run = p.add_run(line)
                run.bold = bold
        elif kind == "table":
            # Table needs a paragraph anchor before it (python-docx requirement).
            if not first:
                cell.add_paragraph()
            _render_md_table(cell, payload, bold=bold)
            first = False
    if first:
        # No blocks at all (empty text); keep the single empty paragraph.
        return text
    return text


# ===========================================================================
# Fill logic
# ===========================================================================
@dataclass
class FillResult:
    filled_fields: List[str] = field(default_factory=list)
    skipped_fields: List[str] = field(default_factory=list)
    figures_embedded: List[int] = field(default_factory=list)
    output_path: str = ""


def fill_template(
    template_id: str,
    content: Dict[str, str],
    output_path: str,
    skill_root: Optional[pathlib.Path] = None,
    figures: Optional[Dict[int, str]] = None,
    figures_width_inches: float = 5.5,
    no_figures: bool = False,
) -> FillResult:
    """Open template .docx, fill cells per TEMPLATES[template_id], save.

    `figures` (optional): mapping {figure_number: image_file_path}. When
    provided, the figure is appended inline at the END of the `details`
    section cell (each figure as a centered paragraph + the PNG), after all
    text/tables have been written. Figures whose path does not exist are
    skipped with a warning. `figures_width_inches` controls the rendered
    width (default 5.5" fits A4 content width).

    Figure auto-discovery (DEFAULT ON): if `figures` is None AND
    `no_figures` is False, this function auto-discovers figures from the
    standard Phase 3 output directory `<skill_root>/../04-diagrams/` (and
    honors the `PATENT_FIGURES_DIR` env var as an override). This makes
    embedded figures the default behavior of the fill pipeline, so users
    no longer have to remember to pass `--figures-dir`. Pass
    `no_figures=True` (or `--no-figures` on the CLI) to opt out.
    """
    if template_id not in TEMPLATES:
        raise ValueError(
            f"Unknown template '{template_id}'. Registered: {list(TEMPLATES)}"
        )
    cfg = TEMPLATES[template_id]

    # Resolve docx path
    if skill_root is None:
        skill_root = _default_skill_root()
    docx_path = (pathlib.Path(skill_root) / cfg["docx_path"])
    if not docx_path.exists():
        raise FileNotFoundError(f"Template .docx not found: {docx_path}")

    # Figure auto-discovery (default ON unless caller passes no_figures=True)
    if figures is None and not no_figures:
        figures = _discover_default_figures_dir(skill_root)
        if figures:
            print(f"  [auto-figures] discovered {len(figures)} figure(s) "
                  f"from default dir; pass --no-figures to disable.")
        # If no default dir / no figures found, treat as 'no figures' silently.

    doc = Document(str(docx_path))
    if len(doc.tables) <= cfg["table_idx"]:
        raise RuntimeError(
            f"Template has {len(doc.tables)} tables, expected index "
            f"{cfg['table_idx']}. Run `inspect` to verify."
        )
    table = doc.tables[cfg["table_idx"]]

    result = FillResult(output_path=output_path)
    figures_embedded: List[int] = []
    # Table whitelist: terminology (Section 7), references (Section 8), and
    # details (Section 4 技术方案的详细阐述) may contain native Word tables
    # in the rendered .docx. Other content sections (background / prior art
    # / problems / invention points / effects / alternatives) are kept as
    # plain text — any markdown pipe in their value stays as a literal '|'
    # character rather than being rendered as a nested table. Section 4 is
    # explicitly allowed because it carries decision-tree tables, Hard Gate
    # tables, and parameter tables that genuinely need 2D layout.
    _TABLE_ALLOWED_FIELDS = {"terminology", "references", "details"}
    for field_name, (row, col) in cfg["fields"].items():
        if field_name not in content:
            result.skipped_fields.append(field_name)
            continue
        value = content[field_name]
        # Special handling for `published` boolean-like
        if field_name == "published":
            value = "是" if str(value).lower() in ("yes", "true", "1", "是") else "否"
        if row >= len(table.rows):
            result.skipped_fields.append(f"{field_name} (row {row} out of range)")
            continue
        cells = table.rows[row].cells
        if col >= len(cells):
            result.skipped_fields.append(f"{field_name} (col {col} out of range)")
            continue
        # Pass the raw value (not str(value)) so set_cell_text -> 
        # _coerce_structured_to_markdown can convert list-of-lists /
        # list-of-dicts into a markdown table before rendering.
        # Stripping/str() happens inside set_cell_text.
        set_cell_text(
            cells[col],
            value,
            allow_tables=(field_name in _TABLE_ALLOWED_FIELDS),
        )
        result.filled_fields.append(field_name)
        # Inline figure embedding: append figures to the `details` cell once
        # its text/tables are written. Each figure is a centered paragraph
        # (caption + image). Missing image files are reported, not fatal.
        if field_name == "details" and figures:
            embedded_now = _embed_figures_in_cell(
                cells[col], figures, figures_width_inches
            )
            figures_embedded.extend(embedded_now)

    result.figures_embedded = figures_embedded
    doc.save(output_path)
    return result


def _embed_figures_in_cell(
    cell, figures: Dict[int, str], width_inches: float
) -> List[int]:
    """Append all figures (sorted by figure number) inline at the end of `cell`.

    Each figure renders as: a centered caption paragraph "图 N" followed by a
    centered paragraph holding the PNG at `width_inches`. Returns the list of
    figure numbers actually embedded (skipping missing files).
    """
    embedded: List[int] = []
    # Spacing paragraph to separate figures from preceding text/table.
    cell.add_paragraph()
    for fig_num in sorted(figures.keys()):
        img_path = figures[fig_num]
        if not os.path.exists(img_path):
            print(f"  [WARN] figure {fig_num} missing: {img_path}")
            continue
        # Caption paragraph (centered, bold).
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"图 {fig_num}")
        cap_run.bold = True
        cap_run.font.size = Pt(10)
        # Image paragraph (centered).
        pic_para = cell.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic_para.add_run()
        try:
            pic_run.add_picture(img_path, width=Inches(width_inches))
            embedded.append(fig_num)
        except Exception as e:
            print(f"  [WARN] figure {fig_num} embed failed: {e}")
            # Remove the now-empty picture paragraph to avoid blank space.
            cell._tc.remove(pic_para._p)
    return embedded


def _default_skill_root() -> pathlib.Path:
    # scripts/ is one level below skill root
    return pathlib.Path(__file__).resolve().parent.parent


def _discover_default_figures_dir(skill_root: pathlib.Path) -> Dict[int, str]:
    """Auto-discover figures from the standard patent-forge Phase 3 output.

    Resolution order (first match wins):
    1. `PATENT_FIGURES_DIR` env var (absolute path).
    2. `<skill_root>/../04-diagrams/` — the canonical Phase 3 figures
       sub-directory used by the patent-forge workflow. If the disclosure
       project is in a sibling folder (e.g. `patent-forge-output/<case>/
       04-diagrams/`), the caller should pass `--figures-dir` explicitly,
       OR set `PATENT_FIGURES_DIR`.

    Returns `{}` if no dir is found or it contains no `fig<N>_*.png` files.
    Never raises (figures are best-effort).
    """
    env_dir = os.environ.get("PATENT_FIGURES_DIR")
    candidates: List[pathlib.Path] = []
    if env_dir:
        candidates.append(pathlib.Path(env_dir))
    candidates.append(pathlib.Path(skill_root).parent / "04-diagrams")
    for d in candidates:
        try:
            mapping = _discover_figures(str(d))
        except FileNotFoundError:
            continue
        if mapping:
            return mapping
    return {}


def _discover_figures(figures_dir: str) -> Dict[int, str]:
    """Scan `figures_dir` for files matching `fig<N>_*.png` and return a
    mapping {N: absolute_path}, sorted by N. SVG companions are ignored
    (Word cannot embed SVG inline; PNG is always preferred for embedding).
    Files not matching the pattern are ignored. Raises if the directory
    does not exist.
    """
    d = pathlib.Path(figures_dir)
    if not d.is_dir():
        raise FileNotFoundError(f"--figures-dir not found: {figures_dir}")
    mapping: Dict[int, str] = {}
    # fig1_..., fig2_... up to fig99. Capture the leading integer.
    pat = re.compile(r"^fig(\d+)_.*\.png$", re.IGNORECASE)
    for p in sorted(d.iterdir()):
        m = pat.match(p.name)
        if m:
            mapping[int(m.group(1))] = str(p.resolve())
    return mapping


# ===========================================================================
# Inspect logic — for adding new templates
# ===========================================================================
def inspect_template(docx_path: str) -> None:
    """Print every table's structure to help derive field->(row,col) mapping."""
    doc = Document(docx_path)
    print(f"File: {docx_path}")
    print(f"Tables: {len(doc.tables)}")
    for ti, t in enumerate(doc.tables):
        print(f"\n=== Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols ===")
        for ri, row in enumerate(t.rows):
            seen_tc = []
            cells_info = []
            for ci, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    tag = "<merged>"
                else:
                    seen_tc.append(tc_id)
                    txt = cell.text.strip().replace("\n", " | ")
                    tag = (txt[:60] + "…") if len(txt) > 60 else txt
                cells_info.append(f"[{ci}]{tag!r}")
            print(f"  Row {ri:2d}: {' '.join(cells_info)}")


# ===========================================================================
# Built-in sample content (AR gesture) — used when --content is omitted
# ===========================================================================
SAMPLE_CONTENT = {
    "case_name":    "一种基于手势识别的增强现实交互方法及系统",
    "inventors":    "张三、李四",
    "applicant":    "XX科技有限公司",
    "tech_contact": "王五",
    "tech_phone":   "138-0000-0000",
    "tech_email":   "wangwu@example.com",
    "published":    "no",
    "bg_field": (
        "本发明涉及增强现实（Augmented Reality, AR）与人机交互技术领域，"
        "特别涉及一种基于手势识别的 AR 交互方法及系统。\n"
        "增强现实技术将虚拟信息叠加至真实世界，已在工业维修、教育培训、医疗辅助等领域得到广泛应用。"
        "手势作为一种自然、直观的交互方式，被认为是 AR 场景中最具潜力的输入模态之一。"
        "然而，现有 AR 手势交互方案存在以下不足：传统基于外观的方法在复杂光照和遮挡场景下识别精度显著下降；"
        "现有轻量化模型难以同时兼顾精度与实时性，端侧推理延迟普遍超过 80ms；"
        "手势指令映射逻辑僵化，缺乏基于上下文的自适应能力。"
    ),
    "closest_prior": (
        "经全球专利库（EPO DOCDB/INPADOC，覆盖 100+ 国家）检索，与本方案最接近的现有技术为：\n"
        "（1）CN114440000A（2022）——一种基于手势的 AR 标注方法。该专利通过肤色分割与模板匹配识别预设手势，成本低但精度受光照影响大。\n"
        "（2）US20210373015A1（2021）——Wearable AR gesture control system。采用 MediaPipe 端侧推理，延迟约 50ms。\n"
        "注：以\"AR gesture recognition real-time end-side\"为关键词在全球专利库中未检索到直接相关专利。"
    ),
    "problems": (
        "现有技术存在以下不足：\n"
        "（1）缺乏一种在端侧设备上同时兼顾高精度（>95%）与低延迟（<30ms）的手部关键点检测网络结构；\n"
        "（2）缺乏一种基于时序卷积网络（TCN）的动态手势序列建模方法；\n"
        "（3）缺乏一种基于场景上下文的自适应手势指令映射机制。"
    ),
    "invention_pts": (
        "本发明提出一种基于手势识别的 AR 交互方法及系统，核心技术方案包括三个层面：\n"
        "（1）关键点检测层面：提出基于 MobileNetV3 主干并经通道剪枝的轻量化手部关键点检测网络。\n"
        "（2）动态手势识别层面：提出基于时序卷积网络（TCN）的序列建模方法。\n"
        "（3）指令映射层面：提出基于场景上下文的自适应映射策略。"
    ),
    "details": (
        "4.1 系统整体架构\n本方案所述 AR 交互系统包括：图像采集模块、关键点检测模块、手势识别模块、指令映射模块、AR 渲染模块。\n"
        "【图 1】基于手势识别的 AR 交互系统整体架构示意图\n\n"
        "4.2 轻量化手部关键点检测网络\n主干采用 MobileNetV3-Large，输入 224×224×3，输出 21 个手部关键点坐标。通道剪枝比例 0.5 ≤ r ≤ 0.75。\n"
        "【图 2】轻量化手部关键点检测网络结构示意图\n\n"
        "4.3 基于 TCN 的动态手势识别\n步骤一：从连续 T=30 帧关键点序列中提取时空特征；\n步骤二：4 层时序卷积（kernel=3，dilation=[1,2,4,8]）建模长程依赖；\n步骤三：经全连接层输出 N 类动态手势概率分布。\n"
        "【图 3】动态手势识别方法流程图\n\n"
        "4.4 基于场景上下文的自适应指令映射\n指令映射模块维护\"空闲/选择/操作/确认\"四态状态机，同一手势在不同状态下映射为不同交互指令。\n"
        "【图 4】指令自适应映射机制示意图"
    ),
    "effects": (
        "（1）通过轻量化手部关键点检测网络，在保证识别精度≥95% 的前提下，将端侧推理延迟降至 30ms 以内。\n"
        "（2）通过基于 TCN 的动态手势序列建模，相比传统 3D CNN 计算量降低约 60%。\n"
        "（3）通过场景上下文自适应映射机制，使同一手势在不同场景下产生不同交互效果，显著提升了交互自然性。"
    ),
    "alternatives": (
        "替代方案一：采用 Transformer 替代 TCN 进行动态手势建模，精度提升约 2%，但计算量增加 20%。\n"
        "替代方案二：采用 RGB-D 融合输入替代纯深度图输入，弱光环境下精度提升 3-5%。"
    ),
    "terminology": (
        "术语/缩略语 | 解释说明\n"
        "AR | Augmented Reality，增强现实\n"
        "CNN | Convolutional Neural Network，卷积神经网络\n"
        "TCN | Temporal Convolutional Network，时序卷积网络\n"
        "MediaPipe | Google 开源的跨平台 ML 推理框架，含手部关键点检测模型\n"
        "MobileNetV3 | Google 提出的轻量化卷积神经网络架构\n"
        "RGB-D | Red Green Blue-Depth，彩色图+深度图四通道输入"
    ),
    "references": (
        "[1] CN114440000A (2022) — 一种基于手势的 AR 标注方法。\n"
        "[2] US20210373015A1 (2021) — Wearable AR gesture control system。\n"
        "[3] Zhang et al., \"MobileNetV3: Searching for MobileNetV3\", ICCV 2019。\n"
        "[4] Bai et al., \"An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling\", arXiv 2018。"
    ),
}


# ===========================================================================
# CLI
# ===========================================================================
def main():
    parser = argparse.ArgumentParser(
        prog="fill_acip_template",
        description="Fill patent disclosure templates with content",
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    # fill
    p_fill = sub.add_parser("fill", help="Fill a registered template")
    p_fill.add_argument("--template", default="acip", choices=list(TEMPLATES))
    p_fill.add_argument("--content", help="Path to JSON file with field values")
    p_fill.add_argument("--output", required=True, help="Output .docx path")
    p_fill.add_argument("--skill-root", help="Override skill root directory")
    p_fill.add_argument(
        "--figures-dir",
        help=(
            "Directory containing figure image files (PNG). Figures named "
            "fig<N>_*.png (e.g. fig1_system_architecture.png) are embedded "
            "inline at the end of the 'details' section, ordered by N. "
            "If omitted, figures are auto-discovered from the default "
            "Phase 3 output dir (see --no-figures)."
        ),
    )
    p_fill.add_argument(
        "--figure-width",
        type=float,
        default=5.5,
        help="Rendered figure width in inches (default 5.5, A4 content width).",
    )
    p_fill.add_argument(
        "--no-figures",
        action="store_true",
        help=(
            "Disable default figure auto-discovery. Use this when you want a "
            "text-only .docx (figures added later by hand)."
        ),
    )

    # inspect
    p_ins = sub.add_parser("inspect", help="Inspect a template's table layout")
    p_ins.add_argument("--docx", required=True, help="Template .docx to inspect")

    # list
    sub.add_parser("list", help="List registered templates")

    args = parser.parse_args()

    if args.cmd == "list":
        print("Registered templates:")
        for tid, cfg in TEMPLATES.items():
            n_fields = len(cfg["fields"])
            print(f"  {tid:10s}  docx={cfg['docx_path']}  "
                  f"table_idx={cfg['table_idx']}  fields={n_fields}")
        return

    if args.cmd == "inspect":
        inspect_template(args.docx)
        return

    if args.cmd == "fill":
        content = SAMPLE_CONTENT
        if args.content:
            with open(args.content, "r", encoding="utf-8") as f:
                content = json.load(f)
        # Figure source: explicit --figures-dir overrides everything.
        # Otherwise leave `figures=None` and let fill_template do default
        # auto-discovery (honoring PATENT_FIGURES_DIR + <skill_root>/../04-diagrams).
        # --no-figures short-circuits the whole figure pipeline.
        figures = _discover_figures(args.figures_dir) if args.figures_dir else None
        result = fill_template(
            args.template,
            content,
            args.output,
            args.skill_root,
            figures=figures,
            figures_width_inches=args.figure_width,
            no_figures=args.no_figures,
        )
        print(f"[OK] Saved: {result.output_path}")
        print(f"  Filled ({len(result.filled_fields)}): {result.filled_fields}")
        if result.skipped_fields:
            print(f"  Skipped ({len(result.skipped_fields)}): {result.skipped_fields}")
        if result.figures_embedded:
            print(f"  Figures embedded: {result.figures_embedded}")
        elif not args.no_figures:
            print("  Figures embedded: (none found — no default dir matched)")


if __name__ == "__main__":
    main()
