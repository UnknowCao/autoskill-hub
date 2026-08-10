"""Generate synthetic test samples for x2md validation.
Creates three test files in tests/samples/:
  1. test_formula.docx  — inline math with $...$ (tests formula escaping fix)
  2. test_merged.docx   — table with rowspan/colspan (tests table detection)
  3. test_formula.xlsx  — formula cells =A2+B2 etc. (tests XLSX formula eval)
"""
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "samples")
os.makedirs(OUT_DIR, exist_ok=True)

# ── 1. test_formula.docx ──────────────────────────────────────────────
from docx import Document

doc = Document()
doc.add_heading("Formula Escaping Test", level=1)
doc.add_paragraph(
    "This document contains inline math formulas that markitdown 0.1.7 "
    "incorrectly escapes. After conversion, $...$ must NOT contain \\* \\_ \\^."
)
doc.add_heading("Simple formulas", level=2)
doc.add_paragraph("Multiplication: $a * b = c$")
doc.add_paragraph("Subscript: $x_1 + x_2 = x_3$")
doc.add_paragraph("Superscript: $e^{i\\pi} + 1 = 0$")
doc.add_paragraph("Combined: $x_1^2 + x_2^2 = r^2$")
doc.add_paragraph("Matrix-like: $a_{11} * b_{1} + a_{12} * b_{2}$")
doc.add_heading("Formulas in context", level=2)
doc.add_paragraph(
    "The energy is given by $E = m * c^2$, where $m$ is mass and $c$ is "
    "the speed of light. For a system with $n$ particles, the total energy "
    "is $E_{total} = \\sum_{i=1}^{n} m_i * c^2$."
)
doc.add_paragraph(
    "The quadratic formula $x = \\frac{-b \\pm \\sqrt{b^2 - 4ac}}{2a}$ "
    "is used when $a * x^2 + b * x + c = 0$ and $a \\neq 0$."
)
out = os.path.join(OUT_DIR, "test_formula.docx")
doc.save(out)
print(f"[OK] {out}")

# ── 2. test_merged.docx ───────────────────────────────────────────────
doc = Document()
doc.add_heading("Table Merge Test", level=1)
doc.add_paragraph(
    "This document contains a table with rowspan and colspan. "
    "markitdown flattens these incorrectly — columns shift left."
)

doc.add_heading("2×2 with rowspan", level=2)
table = doc.add_table(rows=3, cols=2, style="Table Grid")
# Header
table.cell(0, 0).text = "Item"
table.cell(0, 1).text = "Value"
# Row 1: Item A spans 2 rows, Value = 100
a = table.cell(1, 0)
a.text = "A (rowspan=2)"
b = table.cell(1, 1)
b.text = "100"
# Row 2: Value = 200 (Item merged from above)
table.cell(2, 1).text = "200"
# Merge Item A vertically
a.merge(table.cell(2, 0))

doc.add_heading("3×3 with colspan header", level=2)
table2 = doc.add_table(rows=4, cols=3, style="Table Grid")
table2.cell(0, 0).text = "Header spanning 2 cols"
table2.cell(0, 0).merge(table2.cell(0, 1))
table2.cell(0, 2).text = "Col C"
table2.cell(1, 0).text = "A1"
table2.cell(1, 1).text = "B1"
table2.cell(1, 2).text = "C1"
table2.cell(2, 0).text = "A2"
table2.cell(2, 1).text = "B2"
table2.cell(2, 2).text = "C2"
table2.cell(3, 0).text = "A3"
table2.cell(3, 1).text = "B3"
table2.cell(3, 2).text = "C3"

doc.add_heading("Nested structure (semantic test)", level=2)
doc.add_paragraph(
    "Below is a table where one cell conceptually contains a sub-structure. "
    "This tests the nested_table detection path."
)
table3 = doc.add_table(rows=3, cols=2, style="Table Grid")
table3.cell(0, 0).text = "Parameter"
table3.cell(0, 1).text = "Specification"
table3.cell(1, 0).text = "Voltage Range"
table3.cell(1, 1).text = "3.3V ± 5%\n- Min: 3.135V\n- Typ: 3.300V\n- Max: 3.465V"
table3.cell(2, 0).text = "Temperature"
table3.cell(2, 1).text = "-40°C to +85°C\n- Storage: -55°C to +125°C"

out = os.path.join(OUT_DIR, "test_merged.docx")
doc.save(out)
print(f"[OK] {out}")

# ── 3. test_formula.xlsx ──────────────────────────────────────────────
from openpyxl import Workbook

wb = Workbook()
ws = wb.active
ws.title = "Budget"

# Header
ws["A1"] = "Item"
ws["B1"] = "Qty"
ws["C1"] = "Unit Price"
ws["D1"] = "Subtotal"
ws["E1"] = "Tax (13%)"
ws["F1"] = "Total"

# Data rows with formulas
data = [
    ("Widget A", 10, 25.50),
    ("Widget B", 5, 100.00),
    ("Gadget C", 20, 7.75),
    ("Gadget D", 3, 250.00),
]
for i, (item, qty, price) in enumerate(data, start=2):
    ws.cell(row=i, column=1, value=item)
    ws.cell(row=i, column=2, value=qty)
    ws.cell(row=i, column=3, value=price)
    # Subtotal: =B*C  (NO cached value — this is the NaN bug trigger)
    ws.cell(row=i, column=4).value = f"=B{i}*C{i}"
    # Tax: =D*13%
    ws.cell(row=i, column=5).value = f"=D{i}*0.13"
    # Total: =D+E
    ws.cell(row=i, column=6).value = f"=D{i}+E{i}"

# Summary row with SUM
summary_row = len(data) + 2
ws.cell(row=summary_row, column=1, value="TOTAL")
ws.cell(row=summary_row, column=4).value = f"=SUM(D2:D{summary_row-1})"
ws.cell(row=summary_row, column=5).value = f"=SUM(E2:E{summary_row-1})"
ws.cell(row=summary_row, column=6).value = f"=SUM(F2:F{summary_row-1})"

# No cached values written → all formula cells show NaN in base markitdown
out = os.path.join(OUT_DIR, "test_formula.xlsx")
wb.save(out)
print(f"[OK] {out}")
print(f"\nAll 3 synthetic test samples generated in {OUT_DIR}")
